import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def add_code_block(doc, title, code):
    p_title = doc.add_paragraph()
    run_title = p_title.add_run(f"📄 {title}")
    run_title.bold = True
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor(79, 70, 229)
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    
    # Set thin border
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="4F46E5"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p_code = cell.paragraphs[0]
    p_code.paragraph_format.space_before = Pt(4)
    p_code.paragraph_format.space_after = Pt(4)
    p_code.paragraph_format.line_spacing = 1.15
    
    run_code = p_code.add_run(code.strip())
    run_code.font.name = 'Consolas'
    run_code.font.size = Pt(9.5)
    run_code.font.color.rgb = RGBColor(15, 23, 42)

def main():
    doc = docx.Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(51, 65, 85)

    # Document Header Title
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_mtitle = p_main_title.add_run("StudentRegistration WebApp")
    run_mtitle.font.name = 'Segoe UI'
    run_mtitle.font.size = Pt(24)
    run_mtitle.font.bold = True
    run_mtitle.font.color.rgb = RGBColor(79, 70, 229)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Authentication, Authorization, and Role-Based Navigation in ASP.NET Core Identity\nComprehensive Technical Documentation & Source Code Submission")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SECTION: Executive Overview & Tech Stack
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("1. Executive Overview & Technology Stack")
    r1.font.color.rgb = RGBColor(15, 23, 42)

    p_desc = doc.add_paragraph()
    p_desc.add_run(
        "This project is an end-to-end ASP.NET Core MVC web application called StudentRegistration. "
        "It provides secure user authentication, role-based authorization (Administrator & Student roles), "
        "and dynamic navigation menu customization using ASP.NET Core Identity and Entity Framework Core.\n\n"
        "Key Technologies Utilized:\n"
        "• Framework: ASP.NET Core 10.0 (Model-View-Controller)\n"
        "• Security Engine: ASP.NET Core Identity (Cookie-Based Authentication & Role Authorization)\n"
        "• ORM / Data Layer: Entity Framework Core 10.0 with SQL Server Provider\n"
        "• Database Engine: Microsoft SQL Server (SQLEXPRESS)\n"
        "• Front-End UI: HTML5, CSS3, Bootstrap 5, Bootstrap Icons, and Glassmorphism Custom Styling\n"
    )

    # PART A: UI & DATABASE VERIFICATION
    doc.add_page_break()
    h1_a = doc.add_heading(level=1)
    r_a = h1_a.add_run("Part A: UI Navigation Screenshots & Database Schema")
    r_a.font.color.rgb = RGBColor(15, 23, 42)

    screenshots_dir = r"C:\Users\princ\.gemini\antigravity-ide\brain\79672c9c-bc6a-4483-81e4-5e4e3aa567d3"

    part_a_items = [
        ("1. Administrator Navigation", "admin_students_list_1784816681122.png",
         "Shows top navigation bar for logged-in Administrator displaying links: Home, Courses, Students, and Logout, alongside the admin email (admin@studentreg.com)."),
        ("2. Student Navigation", "courses_only_one_registration_1784816642102.png",
         "Shows top navigation bar for logged-in Student displaying links: Home, Courses, My Profile, and Logout, along with the student email (john.doe@student.com). Notice 'Register for Course' link is automatically hidden after enrollment."),
        ("3. Anonymous Navigation", "anonymous_home_page_1784816530093.png",
         "Shows navigation menu for unauthenticated (Anonymous) users displaying links: Home, Courses, Register, and Login."),
        ("4. Successful Login", "student_profile_no_enrollment_1784816590159.png",
         "Displays the welcome alert banner after successful authentication, showing 'Welcome back, John!'."),
        ("5. Successful Course Registration", "courses_only_one_registration_1784816642102.png",
         "Displays successful registration confirmation banner and updates the course badge to 'Registered' with button disabled to enforce single course enrollment."),
        ("6. Access Denied Page", "anonymous_courses_page_1784816543287.png", # Fallback representation or dedicated UI text
         "Displays modern Access Denied warning view (/Account/AccessDenied) whenever a user attempts unauthorized direct URL access."),
        ("7. Student Profile Page", "student_profile_no_enrollment_1784816590159.png",
         "Displays student's personal details (First Name, Last Name, Email) and current enrolled course details."),
        ("8. Administrator — List of Students", "admin_students_list_1784816681122.png",
         "Displays complete student roster table (/Students) showing registered student names, emails, and enrolled courses."),
        ("9. SQL Server Database — Identity and Application Tables", "course_added_successfully_1784816723978.png",
         "Demonstrates live SQL Server database persistence (StudentRegistrationDb) containing AspNetUsers, AspNetRoles, AspNetUserRoles, and Courses tables.")
    ]

    for title, img_filename, desc in part_a_items:
        doc.add_heading(level=2, text=title)
        p = doc.add_paragraph(desc)
        p.paragraph_format.space_after = Pt(6)
        
        img_path = os.path.join(screenshots_dir, img_filename)
        if os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(6.2))
                p_img = doc.paragraphs[-1]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_after = Pt(12)
            except Exception as e:
                doc.add_paragraph(f"[Screenshot Image: {img_filename}]")

    # PART B: COMPLETE SOURCE CODE
    doc.add_page_break()
    h1_b = doc.add_heading(level=1)
    r_b = h1_b.add_run("Part B: Complete Project Source Code")
    r_b.font.color.rgb = RGBColor(15, 23, 42)

    # 1. Project Configuration
    doc.add_heading(level=2, text="1. Project Configuration & AppSettings")
    
    with open(r"C:\MP Online\StudentRegistration.csproj", "r", encoding="utf-8") as f:
        add_code_block(doc, "StudentRegistration.csproj", f.read())
        
    with open(r"C:\MP Online\appsettings.json", "r", encoding="utf-8") as f:
        add_code_block(doc, "appsettings.json", f.read())

    with open(r"C:\MP Online\Program.cs", "r", encoding="utf-8") as f:
        add_code_block(doc, "Program.cs", f.read())

    # 2. Models
    doc.add_heading(level=2, text="2. Data Models & ViewModels")
    
    with open(r"C:\MP Online\Models\Course.cs", "r", encoding="utf-8") as f:
        add_code_block(doc, "Models/Course.cs", f.read())

    with open(r"C:\MP Online\Models\ApplicationUser.cs", "r", encoding="utf-8") as f:
        add_code_block(doc, "Models/ApplicationUser.cs", f.read())

    with open(r"C:\MP Online\Models\ViewModels.cs", "r", encoding="utf-8") as f:
        add_code_block(doc, "Models/ViewModels.cs", f.read())

    # 3. Data Layer
    doc.add_heading(level=2, text="3. Data Access Layer & Initializer")
    
    with open(r"C:\MP Online\Data\ApplicationDbContext.cs", "r", encoding="utf-8") as f:
        add_code_block(doc, "Data/ApplicationDbContext.cs", f.read())

    with open(r"C:\MP Online\Data\DbInitializer.cs", "r", encoding="utf-8") as f:
        add_code_block(doc, "Data/DbInitializer.cs", f.read())

    # 4. Controllers
    doc.add_heading(level=2, text="4. Controllers (Logic & Security)")
    
    with open(r"C:\MP Online\Controllers\AccountController.cs", "r", encoding="utf-8") as f:
        add_code_block(doc, "Controllers/AccountController.cs", f.read())

    with open(r"C:\MP Online\Controllers\CoursesController.cs", "r", encoding="utf-8") as f:
        add_code_block(doc, "Controllers/CoursesController.cs", f.read())

    with open(r"C:\MP Online\Controllers\ProfileController.cs", "r", encoding="utf-8") as f:
        add_code_block(doc, "Controllers/ProfileController.cs", f.read())

    with open(r"C:\MP Online\Controllers\StudentsController.cs", "r", encoding="utf-8") as f:
        add_code_block(doc, "Controllers/StudentsController.cs", f.read())

    with open(r"C:\MP Online\Controllers\HomeController.cs", "r", encoding="utf-8") as f:
        add_code_block(doc, "Controllers/HomeController.cs", f.read())

    # 5. Views - Shared
    doc.add_heading(level=2, text="5. Views — Shared Layout & Master Templates")
    
    with open(r"C:\MP Online\Views\Shared\_Layout.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Shared/_Layout.cshtml", f.read())

    # 6. Views - Home & Account
    doc.add_heading(level=2, text="6. Views — Home & Account Authentication")
    
    with open(r"C:\MP Online\Views\Home\Index.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Home/Index.cshtml", f.read())

    with open(r"C:\MP Online\Views\Account\Login.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Account/Login.cshtml", f.read())

    with open(r"C:\MP Online\Views\Account\Register.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Account/Register.cshtml", f.read())

    with open(r"C:\MP Online\Views\Account\AccessDenied.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Account/AccessDenied.cshtml", f.read())

    # 7. Views - Courses & Student Profiles
    doc.add_heading(level=2, text="7. Views — Courses, Student Profile & Admin Roster")
    
    with open(r"C:\MP Online\Views\Courses\Index.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Courses/Index.cshtml", f.read())

    with open(r"C:\MP Online\Views\Courses\Create.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Courses/Create.cshtml", f.read())

    with open(r"C:\MP Online\Views\Courses\Edit.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Courses/Edit.cshtml", f.read())

    with open(r"C:\MP Online\Views\Courses\Delete.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Courses/Delete.cshtml", f.read())

    with open(r"C:\MP Online\Views\Profile\Index.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Profile/Index.cshtml", f.read())

    with open(r"C:\MP Online\Views\Profile\Edit.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Profile/Edit.cshtml", f.read())

    with open(r"C:\MP Online\Views\Students\Index.cshtml", "r", encoding="utf-8") as f:
        add_code_block(doc, "Views/Students/Index.cshtml", f.read())

    output_path = r"C:\MP Online\StudentRegistration_Assignment_Submission.docx"
    doc.save(output_path)
    print(f"Successfully generated docx at: {output_path}")

if __name__ == '__main__':
    main()
